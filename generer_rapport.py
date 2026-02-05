"""
Script de generation du Rapport de Mise en Place -- PSL-CFX
Genere un document Word (.docx) avec la section Logistique remplie
et des placeholders pour les autres categories.

Source : CSV journalier reconstitue (donnees_journalieres_reconstituees.csv)
"""

import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np
import os
import io

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT


# --- Configuration ---
DATA_PATH = "data/logistics/donnees_journalieres_reconstituees.csv"
OUTPUT_FILE = "Rapport_Mise_En_Place_PSL-CFX_v2.docx"
CHARTS_DIR = "charts_temp"

# Couleurs
BLEU = RGBColor(31, 119, 180)
ROUGE = RGBColor(214, 39, 40)
VERT = RGBColor(44, 160, 44)
GRIS = RGBColor(100, 100, 100)
BLEU_FONCE = RGBColor(0, 51, 102)
BLANC = RGBColor(255, 255, 255)

MOIS_LABELS = [
    "Jan", "Fev", "Mar", "Avr", "Mai", "Jun",
    "Jul", "Aou", "Sep", "Oct", "Nov", "Dec",
]

JOURS_LABELS = ["Lun", "Mar", "Mer", "Jeu", "Ven", "Sam", "Dim"]


# --- Fonctions utilitaires ---

def setup_charts_dir():
    os.makedirs(CHARTS_DIR, exist_ok=True)


def save_chart(fig, name):
    path = os.path.join(CHARTS_DIR, f"{name}.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = BLANC
        set_cell_shading(cell, "003366")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EBF5FB")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLEU_FONCE
    return h


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(text)
        for run in p.runs:
            if not run.bold:
                run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p


def add_sub_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet 2")
    run = p.add_run(text)
    run.font.size = Pt(9)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def add_constat_box(doc, constat_text):
    p = doc.add_paragraph()
    run = p.add_run("Constat : ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = ROUGE
    run2 = p.add_run(constat_text)
    run2.font.size = Pt(10)
    return p


# --- Preparation des donnees ---

def load_daily_data():
    """Charge le CSV journalier et retourne le DataFrame brut."""
    df = pd.read_csv(DATA_PATH)
    df["date"] = pd.to_datetime(df["date"])
    df["year"] = df["year"].astype(int)
    df["month"] = df["month"].astype(int)
    return df


def agg_annual(df):
    """Agregation annuelle : somme des valeurs par annee/indicateur/sous-indicateur (tous sites confondus)."""
    agg = (
        df.groupby(["year", "indicateur", "sous_indicateur", "unite"])[["value", "value_crise"]]
        .sum()
        .reset_index()
    )
    agg["ecart"] = agg["value_crise"] - agg["value"]
    agg["variation_pct"] = np.where(agg["value"] == 0, np.nan, (agg["ecart"] / agg["value"]) * 100)
    return agg


def agg_annual_by_site(df):
    """Agregation annuelle par site."""
    return (
        df.groupby(["year", "site_code", "indicateur", "sous_indicateur", "unite"])[["value", "value_crise"]]
        .sum()
        .reset_index()
    )


def agg_monthly(df):
    """Agregation mensuelle (tous sites confondus)."""
    return (
        df.groupby(["year", "month", "indicateur", "sous_indicateur", "unite"])[["value", "value_crise"]]
        .sum()
        .reset_index()
    )


def agg_dow(df):
    """Agregation par jour de semaine (moyenne journaliere)."""
    return (
        df.groupby(["dow", "indicateur", "sous_indicateur", "unite"])[["value", "value_crise"]]
        .mean()
        .reset_index()
    )


# ============================================================
# GENERATION DES GRAPHIQUES
# ============================================================

def gen_chart_restauration_annual(annual):
    """Barres annuelles : nombre total de repas Normal vs Crise."""
    df_repas = annual[
        (annual["indicateur"] == "Restauration") &
        (annual["sous_indicateur"] == "Nombre de Repas")
    ].sort_values("year")

    fig, ax = plt.subplots(figsize=(10, 4.5))
    x = np.arange(len(df_repas))
    w = 0.35
    ax.bar(x - w / 2, df_repas["value"], w, label="Situation normale", color="#2ca02c", edgecolor="white")
    ax.bar(x + w / 2, df_repas["value_crise"], w, label="Crise sanitaire (+70 %)", color="#d62728", edgecolor="white")
    ax.set_xlabel("Annee", fontsize=11)
    ax.set_ylabel("Nombre de repas", fontsize=11)
    ax.set_title("Nombre total de repas par an -- Normal vs Crise sanitaire", fontsize=13, fontweight="bold")
    ax.set_xticks(x)
    ax.set_xticklabels(df_repas["year"].astype(int))
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"{v:,.0f}"))
    ax.legend(fontsize=10)
    ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    return save_chart(fig, "restauration_annual")


def gen_chart_restauration_monthly(monthly):
    """Profil mensuel de la restauration (courbes par annee)."""
    df_m = monthly[
        (monthly["indicateur"] == "Restauration") &
        (monthly["sous_indicateur"] == "Nombre de Repas")
    ].copy()

    fig, ax = plt.subplots(figsize=(10, 4.5))
    colors = plt.cm.Set2(np.linspace(0, 1, df_m["year"].nunique()))
    for i, (yr, grp) in enumerate(df_m.groupby("year")):
        grp = grp.sort_values("month")
        ax.plot(grp["month"], grp["value"], "o-", color=colors[i], linewidth=1.5, label=str(yr))
    ax.set_xlabel("Mois", fontsize=11)
    ax.set_ylabel("Nombre de repas", fontsize=11)
    ax.set_title("Profil mensuel de la restauration (situation normale)", fontsize=13, fontweight="bold")
    ax.set_xticks(range(1, 13))
    ax.set_xticklabels(MOIS_LABELS, fontsize=9)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"{v:,.0f}"))
    ax.legend(fontsize=9, title="Annee")
    ax.grid(alpha=0.3)
    plt.tight_layout()
    return save_chart(fig, "restauration_monthly")


def gen_chart_restauration_site(annual_site):
    """Comparaison PLF vs CFX pour la restauration (2015)."""
    df_s = annual_site[
        (annual_site["indicateur"] == "Restauration") &
        (annual_site["sous_indicateur"] == "Nombre de Repas") &
        (annual_site["year"] == 2015)
    ]

    fig, ax = plt.subplots(figsize=(8, 4))
    sites = ["PLF", "CFX"]
    sites_labels = ["Pitie-Salpetriere", "Charles Foix"]
    x = np.arange(2)
    vals_n = [df_s[df_s["site_code"] == s]["value"].sum() for s in sites]
    vals_c = [df_s[df_s["site_code"] == s]["value_crise"].sum() for s in sites]
    ax.bar(x - 0.2, vals_n, 0.35, label="Normal", color="#1f77b4", edgecolor="white")
    ax.bar(x + 0.2, vals_c, 0.35, label="Crise (+70 %)", color="#d62728", edgecolor="white")
    ax.set_xticks(x)
    ax.set_xticklabels(sites_labels, fontsize=10)
    ax.set_title("Restauration par site (2015) -- Normal vs Crise", fontsize=13, fontweight="bold")
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"{v:,.0f}"))
    ax.legend(fontsize=10)
    ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    return save_chart(fig, "restauration_site")


def gen_chart_dechets(annual):
    """Barres horizontales : types de dechets en 2015."""
    df_d = annual[
        (annual["indicateur"] == "Déchets") &
        (annual["year"] == 2015)
    ].copy()

    labels_courts = []
    for s in df_d["sous_indicateur"].values:
        if "Infectieux" in s:
            labels_courts.append("DASRI")
        elif "Ménagers Assimil" in s or "ménagers" in s.lower():
            labels_courts.append("Déchets ménagers (DAE)")
        elif "assimilés aux ordures" in s.lower():
            labels_courts.append("Ordures ménagères")
        elif "lectriques" in s or "Electriques" in s:
            labels_courts.append("DEEE")
        elif "Cartons" in s:
            labels_courts.append("Cartons")
        elif "Papiers" in s:
            labels_courts.append("Papiers")
        elif "Chimiques" in s:
            labels_courts.append("Chimiques")
        else:
            labels_courts.append(s[:30])

    fig, ax = plt.subplots(figsize=(11, 5))
    y = np.arange(len(labels_courts))
    h = 0.35
    ax.barh(y - h / 2, df_d["value"].values, h, label="Normal", color="#1f77b4", edgecolor="white")
    ax.barh(y + h / 2, df_d["value_crise"].values, h, label="Crise (+70 %)", color="#ff7f0e", edgecolor="white")
    ax.set_yticks(y)
    ax.set_yticklabels(labels_courts, fontsize=9)
    ax.set_xlabel("Tonnes", fontsize=11)
    ax.set_title("Volume de dechets par type (2015) -- Normal vs Crise", fontsize=13, fontweight="bold")
    ax.legend(fontsize=10)
    ax.grid(axis="x", alpha=0.3)
    plt.tight_layout()
    return save_chart(fig, "dechets")


def gen_chart_dechets_monthly(monthly):
    """Profil mensuel des DASRI."""
    df_m = monthly[
        (monthly["indicateur"] == "Déchets") &
        (monthly["sous_indicateur"].str.contains("Infectieux", na=False))
    ].copy()

    fig, ax = plt.subplots(figsize=(10, 4.5))
    colors = plt.cm.Set2(np.linspace(0, 1, df_m["year"].nunique()))
    for i, (yr, grp) in enumerate(df_m.groupby("year")):
        grp = grp.sort_values("month")
        ax.plot(grp["month"], grp["value"], "o-", color=colors[i], linewidth=1.5, label=str(yr))
    ax.set_xlabel("Mois", fontsize=11)
    ax.set_ylabel("Volume DASRI (tonnes)", fontsize=11)
    ax.set_title("Saisonnalite des dechets infectieux (DASRI)", fontsize=13, fontweight="bold")
    ax.set_xticks(range(1, 13))
    ax.set_xticklabels(MOIS_LABELS, fontsize=9)
    ax.legend(fontsize=9, title="Annee")
    ax.grid(alpha=0.3)
    plt.tight_layout()
    return save_chart(fig, "dechets_monthly")


def gen_chart_hygiene(annual):
    """Courbes : Locaux et Espaces verts par annee."""
    df_h = annual[annual["indicateur"] == "Hygiène"].copy()
    locaux = df_h[df_h["sous_indicateur"] == "Locaux"].sort_values("year")
    espaces = df_h[df_h["sous_indicateur"] == "Espaces verts"].sort_values("year")

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 4.5))

    ax1.plot(locaux["year"], locaux["value"], "o-", color="#1f77b4", linewidth=2, label="Normal")
    ax1.plot(locaux["year"], locaux["value_crise"], "s--", color="#d62728", linewidth=2, label="Crise (+70 %)")
    ax1.fill_between(locaux["year"], locaux["value"], locaux["value_crise"], alpha=0.15, color="red")
    ax1.set_title("Locaux (m²)", fontsize=12, fontweight="bold")
    ax1.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"{v:,.0f}"))
    ax1.legend(fontsize=9)
    ax1.grid(alpha=0.3)

    ax2.plot(espaces["year"], espaces["value"], "o-", color="#1f77b4", linewidth=2, label="Normal")
    ax2.plot(espaces["year"], espaces["value_crise"], "s--", color="#d62728", linewidth=2, label="Crise (+70 %)")
    ax2.fill_between(espaces["year"], espaces["value"], espaces["value_crise"], alpha=0.15, color="red")
    ax2.set_title("Espaces verts (m²)", fontsize=12, fontweight="bold")
    ax2.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"{v:,.0f}"))
    ax2.legend(fontsize=9)
    ax2.grid(alpha=0.3)

    fig.suptitle("Hygiene : surfaces a entretenir -- Normal vs Crise", fontsize=13, fontweight="bold", y=1.02)
    plt.tight_layout()
    return save_chart(fig, "hygiene")


def gen_chart_hygiene_monthly(monthly):
    """Profil mensuel des surfaces de locaux a nettoyer."""
    df_m = monthly[
        (monthly["indicateur"] == "Hygiène") &
        (monthly["sous_indicateur"] == "Locaux")
    ].copy()

    fig, ax = plt.subplots(figsize=(10, 4.5))
    colors = plt.cm.Set2(np.linspace(0, 1, df_m["year"].nunique()))
    for i, (yr, grp) in enumerate(df_m.groupby("year")):
        grp = grp.sort_values("month")
        ax.plot(grp["month"], grp["value"], "o-", color=colors[i], linewidth=1.5, label=str(yr))
    ax.set_xlabel("Mois", fontsize=11)
    ax.set_ylabel("Surface nettoyee (m²)", fontsize=11)
    ax.set_title("Profil mensuel -- Entretien des locaux (situation normale)", fontsize=13, fontweight="bold")
    ax.set_xticks(range(1, 13))
    ax.set_xticklabels(MOIS_LABELS, fontsize=9)
    ax.legend(fontsize=9, title="Annee")
    ax.grid(alpha=0.3)
    plt.tight_layout()
    return save_chart(fig, "hygiene_monthly")


def gen_chart_lingerie(annual):
    """Courbe annuelle de la quantite de linge."""
    df_l = annual[
        (annual["indicateur"] == "Lingerie") &
        (annual["sous_indicateur"] == "Quantité de linge")
    ].sort_values("year")

    fig, ax = plt.subplots(figsize=(10, 4.5))
    ax.fill_between(df_l["year"], df_l["value"], alpha=0.3, color="#1f77b4")
    ax.fill_between(df_l["year"], df_l["value"], df_l["value_crise"], alpha=0.3, color="#d62728")
    ax.plot(df_l["year"], df_l["value"], "o-", color="#1f77b4", linewidth=2, label="Normal")
    ax.plot(df_l["year"], df_l["value_crise"], "s--", color="#d62728", linewidth=2, label="Crise (+70 %)")
    ax.set_title("Quantite de linge traite (kg/an)", fontsize=13, fontweight="bold")
    ax.set_xlabel("Annee", fontsize=11)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"{v:,.0f}"))
    ax.legend(fontsize=10)
    ax.grid(alpha=0.3)
    plt.tight_layout()
    return save_chart(fig, "lingerie_annual")


def gen_chart_lingerie_site(annual_site):
    """Repartition par site 2015."""
    df_s = annual_site[
        (annual_site["indicateur"] == "Lingerie") &
        (annual_site["sous_indicateur"] == "Quantité de linge") &
        (annual_site["year"] == 2015)
    ]

    fig, ax = plt.subplots(figsize=(8, 4))
    sites = ["PLF", "CFX"]
    sites_labels = ["Pitie-Salpetriere", "Charles Foix"]
    x = np.arange(2)
    vals_n = [df_s[df_s["site_code"] == s]["value"].sum() for s in sites]
    vals_c = [df_s[df_s["site_code"] == s]["value_crise"].sum() for s in sites]
    ax.bar(x - 0.2, vals_n, 0.35, label="Normal", color="#1f77b4", edgecolor="white")
    ax.bar(x + 0.2, vals_c, 0.35, label="Crise (+70 %)", color="#d62728", edgecolor="white")
    ax.set_xticks(x)
    ax.set_xticklabels(sites_labels, fontsize=10)
    ax.set_title("Lingerie par site (2015) -- Normal vs Crise", fontsize=13, fontweight="bold")
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"{v:,.0f}"))
    ax.legend(fontsize=10)
    ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    return save_chart(fig, "lingerie_site")


def gen_chart_lingerie_monthly(monthly):
    """Profil mensuel du linge."""
    df_m = monthly[
        (monthly["indicateur"] == "Lingerie") &
        (monthly["sous_indicateur"] == "Quantité de linge")
    ].copy()

    fig, ax = plt.subplots(figsize=(10, 4.5))
    colors = plt.cm.Set2(np.linspace(0, 1, df_m["year"].nunique()))
    for i, (yr, grp) in enumerate(df_m.groupby("year")):
        grp = grp.sort_values("month")
        ax.plot(grp["month"], grp["value"], "o-", color=colors[i], linewidth=1.5, label=str(yr))
    ax.set_xlabel("Mois", fontsize=11)
    ax.set_ylabel("Linge traite (kg)", fontsize=11)
    ax.set_title("Profil mensuel -- Lingerie (situation normale)", fontsize=13, fontweight="bold")
    ax.set_xticks(range(1, 13))
    ax.set_xticklabels(MOIS_LABELS, fontsize=9)
    ax.legend(fontsize=9, title="Annee")
    ax.grid(alpha=0.3)
    plt.tight_layout()
    return save_chart(fig, "lingerie_monthly")


def gen_chart_magasin(annual):
    """Barres : references gerees par le magasin en 2015."""
    df_m = annual[
        (annual["indicateur"] == "Magasin") &
        (annual["year"] == 2015)
    ].copy().sort_values("sous_indicateur")

    fig, ax = plt.subplots(figsize=(8, 4.5))
    cats = df_m["sous_indicateur"].values
    x = np.arange(len(cats))
    w = 0.35
    ax.bar(x - w / 2, df_m["value"].values, w, label="Normal", color="#2ca02c", edgecolor="white")
    ax.bar(x + w / 2, df_m["value_crise"].values, w, label="Crise (+70 %)", color="#d62728", edgecolor="white")
    ax.set_xticks(x)
    ax.set_xticklabels([c.replace("Références ", "").capitalize() for c in cats], fontsize=10)
    ax.set_ylabel("Nombre de references")
    ax.set_title("References gerees par le magasin (2015)", fontsize=13, fontweight="bold")
    ax.legend(fontsize=10)
    ax.grid(axis="y", alpha=0.3)

    for i, (n, c) in enumerate(zip(df_m["value"].values, df_m["value_crise"].values)):
        ax.text(i - w / 2, n + max(df_m["value_crise"].values) * 0.01, f"{n:,.0f}", ha="center", va="bottom", fontsize=9)
        ax.text(i + w / 2, c + max(df_m["value_crise"].values) * 0.01, f"{c:,.0f}", ha="center", va="bottom", fontsize=9, color="#d62728")

    plt.tight_layout()
    return save_chart(fig, "magasin")


def gen_chart_magasin_monthly(monthly):
    """Profil mensuel du magasin (toutes references confondues)."""
    df_m = monthly[monthly["indicateur"] == "Magasin"].copy()
    agg = df_m.groupby(["year", "month"])[["value"]].sum().reset_index()

    fig, ax = plt.subplots(figsize=(10, 4.5))
    colors = plt.cm.Set2(np.linspace(0, 1, agg["year"].nunique()))
    for i, (yr, grp) in enumerate(agg.groupby("year")):
        grp = grp.sort_values("month")
        ax.plot(grp["month"], grp["value"], "o-", color=colors[i], linewidth=1.5, label=str(yr))
    ax.set_xlabel("Mois", fontsize=11)
    ax.set_ylabel("Nombre de references", fontsize=11)
    ax.set_title("Profil mensuel -- Approvisionnements magasin (situation normale)", fontsize=13, fontweight="bold")
    ax.set_xticks(range(1, 13))
    ax.set_xticklabels(MOIS_LABELS, fontsize=9)
    ax.legend(fontsize=9, title="Annee")
    ax.grid(alpha=0.3)
    plt.tight_layout()
    return save_chart(fig, "magasin_monthly")


def gen_chart_vaguemestre(annual):
    """Barres : plis et colis en 2015."""
    df_v = annual[(annual["indicateur"] == "Vaguemestre") & (annual["year"] == 2015)].copy()
    plis = df_v[df_v["sous_indicateur"] == "Plis affranchis"]
    colis = df_v[df_v["sous_indicateur"] == "Colis"]

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))

    if not plis.empty:
        ax1.bar(["Normal", "Crise"], [plis["value"].values[0], plis["value_crise"].values[0]],
                color=["#1f77b4", "#d62728"], edgecolor="white")
        ax1.set_title("Plis affranchis/an (2015)", fontsize=11, fontweight="bold")
        ax1.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"{v:,.0f}"))
        ax1.grid(axis="y", alpha=0.3)

    if not colis.empty:
        ax2.bar(["Normal", "Crise"], [colis["value"].values[0], colis["value_crise"].values[0]],
                color=["#1f77b4", "#d62728"], edgecolor="white")
        ax2.set_title("Colis/an (2015)", fontsize=11, fontweight="bold")
        ax2.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"{v:,.0f}"))
        ax2.grid(axis="y", alpha=0.3)

    fig.suptitle("Service courrier -- Normal vs Crise", fontsize=13, fontweight="bold", y=1.02)
    plt.tight_layout()
    return save_chart(fig, "vaguemestre")


def gen_chart_vaguemestre_monthly(monthly):
    """Profil mensuel du courrier."""
    df_m = monthly[
        (monthly["indicateur"] == "Vaguemestre") &
        (monthly["sous_indicateur"] == "Plis affranchis")
    ].copy()

    fig, ax = plt.subplots(figsize=(10, 4.5))
    colors = plt.cm.Set2(np.linspace(0, 1, df_m["year"].nunique()))
    for i, (yr, grp) in enumerate(df_m.groupby("year")):
        grp = grp.sort_values("month")
        ax.plot(grp["month"], grp["value"], "o-", color=colors[i], linewidth=1.5, label=str(yr))
    ax.set_xlabel("Mois", fontsize=11)
    ax.set_ylabel("Plis affranchis", fontsize=11)
    ax.set_title("Profil mensuel -- Plis affranchis (situation normale)", fontsize=13, fontweight="bold")
    ax.set_xticks(range(1, 13))
    ax.set_xticklabels(MOIS_LABELS, fontsize=9)
    ax.legend(fontsize=9, title="Annee")
    ax.grid(alpha=0.3)
    plt.tight_layout()
    return save_chart(fig, "vaguemestre_monthly")


def gen_chart_synthese(annual):
    """Barres horizontales : tous les indicateurs (2015)."""
    df_2015 = annual[annual["year"] == 2015].copy()
    synthese = df_2015.groupby("indicateur").agg({"value": "sum", "value_crise": "sum"}).reset_index()

    fig, ax = plt.subplots(figsize=(10, 5))
    y = np.arange(len(synthese))
    h = 0.35
    ax.barh(y - h / 2, synthese["value"], h, label="Normal", color="#1f77b4", edgecolor="white")
    ax.barh(y + h / 2, synthese["value_crise"], h, label="Crise (+70 %)", color="#d62728", edgecolor="white")
    ax.set_yticks(y)
    ax.set_yticklabels(synthese["indicateur"], fontsize=10)
    ax.set_title("Synthese Logistique -- Tous indicateurs (2015)", fontsize=13, fontweight="bold")
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"{v:,.0f}"))
    ax.legend(fontsize=10)
    ax.grid(axis="x", alpha=0.3)
    plt.tight_layout()
    return save_chart(fig, "synthese")


def gen_chart_comparaison_sites(annual_site):
    """Barres groupees : comparaison PLF vs CFX sur tous les indicateurs (2015)."""
    df_2015 = annual_site[annual_site["year"] == 2015]
    agg = df_2015.groupby(["site_code", "indicateur"])["value"].sum().reset_index()

    indicateurs = sorted(agg["indicateur"].unique())
    plf_vals = [agg[(agg["site_code"] == "PLF") & (agg["indicateur"] == ind)]["value"].sum() for ind in indicateurs]
    cfx_vals = [agg[(agg["site_code"] == "CFX") & (agg["indicateur"] == ind)]["value"].sum() for ind in indicateurs]

    fig, ax = plt.subplots(figsize=(11, 5))
    x = np.arange(len(indicateurs))
    w = 0.35
    ax.barh(x - w / 2, plf_vals, w, label="Pitie-Salpetriere (PLF)", color="#1f77b4", edgecolor="white")
    ax.barh(x + w / 2, cfx_vals, w, label="Charles Foix (CFX)", color="#ff7f0e", edgecolor="white")
    ax.set_yticks(x)
    ax.set_yticklabels(indicateurs, fontsize=10)
    ax.set_title("Comparaison des volumes par site (2015, situation normale)", fontsize=13, fontweight="bold")
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"{v:,.0f}"))
    ax.legend(fontsize=10)
    ax.grid(axis="x", alpha=0.3)
    plt.tight_layout()
    return save_chart(fig, "comparaison_sites")


def gen_chart_dow(dow_data):
    """Barres : profil moyen par jour de semaine (restauration)."""
    df_d = dow_data[
        (dow_data["indicateur"] == "Restauration") &
        (dow_data["sous_indicateur"] == "Nombre de Repas")
    ].sort_values("dow")

    fig, ax = plt.subplots(figsize=(8, 4))
    # dow: 0=Sunday .. 6=Saturday -> reorder to Mon-Sun
    dow_order = [1, 2, 3, 4, 5, 6, 0]
    df_ord = df_d.set_index("dow").reindex(dow_order).reset_index()
    x = np.arange(7)
    ax.bar(x, df_ord["value"], color="#1f77b4", edgecolor="white", label="Normal")
    ax.bar(x, df_ord["value_crise"] - df_ord["value"], bottom=df_ord["value"],
           color="#d62728", edgecolor="white", alpha=0.5, label="Surplus crise")
    ax.set_xticks(x)
    ax.set_xticklabels(JOURS_LABELS, fontsize=10)
    ax.set_ylabel("Repas moyens/jour", fontsize=11)
    ax.set_title("Restauration : profil moyen par jour de semaine", fontsize=13, fontweight="bold")
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"{v:,.0f}"))
    ax.legend(fontsize=10)
    ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    return save_chart(fig, "dow_restauration")


# ============================================================
# CONSTRUCTION DU DOCUMENT
# ============================================================

def build_document():
    setup_charts_dir()

    # Chargement et agregations
    df = load_daily_data()
    annual = agg_annual(df)
    annual_site = agg_annual_by_site(df)
    monthly = agg_monthly(df)
    dow_data = agg_dow(df)

    doc = Document()

    # -- Style par defaut --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # ============================
    # PAGE DE TITRE
    # ============================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Rapport de Mise en Place")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BLEU_FONCE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Analyse des donnees hospitalieres\nPSL-CFX (Pitie-Salpetriere / Charles Foix)")
    run.font.size = Pt(16)
    run.font.color.rgb = GRIS

    doc.add_paragraph()

    obj = doc.add_paragraph()
    obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = obj.add_run("Propositions pour gerer les afflux de patients\net se preparer aux crises sanitaires")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = GRIS

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("Fevrier 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = GRIS

    doc.add_page_break()

    # ============================
    # SOMMAIRE
    # ============================
    add_heading_styled(doc, "Sommaire", level=1)
    doc.add_paragraph()

    sommaire_items = [
        ("1.", "Logistique hospitaliere", "Gestion des flux materiels et des approvisionnements"),
        ("2.", "Activite et services", "Volumes d'activite et organisation des soins (a venir)"),
        ("3.", "Capacite d'accueil", "Lits, equipements et plateaux techniques (a venir)"),
        ("4.", "Finance", "Equilibre budgetaire et investissements (a venir)"),
        ("5.", "Patients", "Profil des patients et parcours de soins (a venir)"),
        ("6.", "Qualite", "Satisfaction et indicateurs qualite (a venir)"),
        ("7.", "Ressources humaines", "Effectifs et organisation du personnel (a venir)"),
        ("8.", "Synthese et recommandations", ""),
    ]

    for num, titre, desc in sommaire_items:
        p = doc.add_paragraph()
        run_num = p.add_run(f"{num} ")
        run_num.bold = True
        run_num.font.size = Pt(12)
        run_num.font.color.rgb = BLEU_FONCE

        run_titre = p.add_run(titre)
        run_titre.bold = True
        run_titre.font.size = Pt(12)

        if desc:
            run_desc = p.add_run(f" -- {desc}")
            run_desc.font.size = Pt(10)
            run_desc.font.color.rgb = GRIS

    doc.add_page_break()

    # ============================
    # 1. LOGISTIQUE
    # ============================
    add_heading_styled(doc, "1. Logistique hospitaliere", level=1)

    add_body(doc,
        "La logistique hospitaliere est un pilier essentiel du bon fonctionnement de l'hopital. "
        "Elle couvre l'ensemble des flux materiels necessaires aux soins et a la vie quotidienne "
        "de l'etablissement : restauration des patients et du personnel, gestion du linge, "
        "approvisionnement en fournitures, traitement des dechets, entretien des locaux et "
        "gestion du courrier."
    )

    add_body(doc,
        "Enjeu principal : En situation de crise sanitaire (epidemie, pandemie, afflux massif "
        "de patients), chacun de ces postes logistiques peut etre soumis a une augmentation "
        "brutale de la demande. L'hopital doit etre en mesure d'anticiper ces pics pour maintenir "
        "la qualite des soins et la securite des patients et du personnel."
    )

    add_body(doc,
        "Les donnees analysees couvrent les deux sites du groupe hospitalier :"
    )
    add_bullet(doc, " Hopital de reference, forte activite MCO (Medecine, Chirurgie, Obstetrique)", bold_prefix="Pitie-Salpetriere (PSL) --")
    add_bullet(doc, " Specialise en geriatrie et soins de suite", bold_prefix="Charles Foix (CFX) --")

    add_body(doc,
        "La simulation de crise sanitaire utilisee dans cette analyse repose sur une augmentation "
        "de +70 % de l'activite sur l'ensemble des postes logistiques, conformement aux scenarios "
        "observes lors des pandemies recentes (COVID-19)."
    )

    add_body(doc,
        "Les donnees journalieres reconstituees permettent d'analyser non seulement les volumes "
        "annuels, mais aussi les profils saisonniers (mensuels) et les variations hebdomadaires, "
        "offrant une vision fine de la charge logistique a anticiper."
    )

    # --- 1.1 RESTAURATION ---
    add_heading_styled(doc, "1.1 Restauration hospitaliere", level=2)

    add_body(doc,
        "La restauration est un service critique : chaque patient hospitalise doit recevoir "
        "ses repas quotidiens. En cas d'afflux de patients, le nombre de repas a preparer "
        "augmente proportionnellement. La cuisine centrale doit etre dimensionnee pour absorber "
        "ces variations."
    )

    # Graphique annuel
    chart_path = gen_chart_restauration_annual(annual)
    doc.add_picture(chart_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Chiffres cles
    r2015 = annual[
        (annual["indicateur"] == "Restauration") &
        (annual["sous_indicateur"] == "Nombre de Repas") &
        (annual["year"] == 2015)
    ].iloc[0]
    rq2015 = annual[
        (annual["indicateur"] == "Restauration") &
        (annual["sous_indicateur"] == "Nombre de Repas Quotidien") &
        (annual["year"] == 2015)
    ]

    ecart_repas = r2015["ecart"]
    add_constat_box(doc,
        f"En situation de crise, le nombre de repas a fournir augmenterait de +70 %, soit environ "
        f"{ecart_repas:,.0f} repas supplementaires par an (donnees 2015). "
        f"En 2015, le total s'eleve a {r2015['value']:,.0f} repas en situation normale "
        f"et {r2015['value_crise']:,.0f} en situation de crise."
    )

    if not rq2015.empty:
        rq = rq2015.iloc[0]
        add_body(doc,
            f"Au quotidien, cela represente le passage de {rq['value']:,.0f} repas/jour "
            f"a environ {rq['value_crise']:,.0f} repas/jour."
        )

    # Graphique mensuel
    add_body(doc,
        "Le profil mensuel ci-dessous montre la saisonnalite de la restauration. Les mois d'ete "
        "(juillet-aout) affichent generalement un volume legerement plus faible en raison des conges, "
        "tandis que les mois de forte activite hospitaliere voient des pics de demande."
    )

    chart_path = gen_chart_restauration_monthly(monthly)
    doc.add_picture(chart_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Graphique par site
    add_body(doc,
        "La repartition entre les deux sites montre un desequilibre lie a la taille respective "
        "des etablissements :"
    )

    chart_path = gen_chart_restauration_site(annual_site)
    doc.add_picture(chart_path, width=Inches(5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Repartition PLF/CFX
    plf_rest = annual_site[
        (annual_site["indicateur"] == "Restauration") &
        (annual_site["sous_indicateur"] == "Nombre de Repas") &
        (annual_site["year"] == 2015) &
        (annual_site["site_code"] == "PLF")
    ]["value"].sum()
    total_rest = r2015["value"]
    pct_plf = (plf_rest / total_rest * 100) if total_rest > 0 else 0

    add_body(doc,
        f"Pitie-Salpetriere assure environ {pct_plf:.0f} % des repas, Charles Foix {100 - pct_plf:.0f} %. "
        f"En crise, envisager un reequilibrage si l'un des sites est plus impacte."
    )

    # Graphique jour de semaine
    add_body(doc,
        "Le profil hebdomadaire illustre la charge moyenne de restauration selon le jour de la semaine :"
    )

    chart_path = gen_chart_dow(dow_data)
    doc.add_picture(chart_path, width=Inches(5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Propositions pour gerer l'afflux :")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLEU_FONCE

    add_bullet(doc, " Prevoir des conventions avec des prestataires de restauration externe "
               "(traiteurs agrees par l'ARS) activables sous 48h. Constituer un stock tampon "
               "de repas longue conservation pour couvrir les 72 premieres heures de crise.",
               bold_prefix="Protocole de montee en charge de la cuisine centrale --")

    add_bullet(doc, " Passer a un menu unique simplifie (au lieu de 2-3 choix) pour accelerer "
               "la production. Privilegier les plats en barquettes individuelles pour limiter "
               "les contaminations croisees.",
               bold_prefix="Adaptation des menus en situation de crise --")

    add_bullet(doc, " Identifier dans le plan de crise les agents mobilisables "
               "(personnel en back-office, stagiaires, prestataires). Prevoir un planning de "
               "rotation specifique en mode crise.",
               bold_prefix="Renforcement des effectifs --")

    add_bullet(doc, " En crise, envisager un "
               "reequilibrage si l'un des sites est plus impacte.",
               bold_prefix="Repartition entre sites --")

    # --- 1.2 DECHETS ---
    doc.add_page_break()
    add_heading_styled(doc, "1.2 Gestion des dechets hospitaliers", level=2)

    add_body(doc,
        "Les dechets hospitaliers sont un enjeu majeur de securite sanitaire. Les Dechets "
        "d'Activites de Soins a Risques Infectieux (DASRI) necessitent un traitement specifique "
        "et reglemente. En situation de crise, l'augmentation de l'activite de soins entraine "
        "mecaniquement une hausse des dechets, notamment infectieux."
    )

    chart_path = gen_chart_dechets(annual)
    doc.add_picture(chart_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    dasri = annual[
        (annual["indicateur"] == "Déchets") &
        (annual["sous_indicateur"].str.contains("Infectieux", na=False)) &
        (annual["year"] == 2015)
    ]
    menagers = annual[
        (annual["indicateur"] == "Déchets") &
        (annual["sous_indicateur"].str.contains("Ménagers Assimil", na=False)) &
        (annual["year"] == 2015)
    ]

    if not dasri.empty and not menagers.empty:
        d = dasri.iloc[0]
        m = menagers.iloc[0]
        add_constat_box(doc,
            f"Les DASRI (dechets infectieux) passeraient de {d['value']:,.0f} tonnes a "
            f"{d['value_crise']:,.0f} tonnes en situation de crise (+{d['ecart']:,.0f} t). "
            f"Les dechets menagers augmenteraient de {m['value']:,.0f} tonnes a "
            f"{m['value_crise']:,.0f} tonnes (+{m['ecart']:,.0f} t)."
        )

    # Profil mensuel DASRI
    add_body(doc,
        "Le profil mensuel des DASRI permet d'identifier les periodes de forte production "
        "et d'adapter les capacites de collecte en consequence :"
    )

    chart_path = gen_chart_dechets_monthly(monthly)
    doc.add_picture(chart_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body(doc,
        "On observe que la production de dechets infectieux suit la saisonnalite de l'activite "
        "hospitaliere, avec des variations pouvant atteindre 15 a 20 % entre les mois les plus "
        "calmes et les plus charges."
    )

    p = doc.add_paragraph()
    run = p.add_run("Propositions :")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLEU_FONCE

    add_bullet(doc, " Installer des conteneurs supplementaires de stockage DASRI sur chaque site "
               "(zone dediee, ventilee, securisee). Prevoir des capacites de stockage froid pour "
               "les dechets anatomiques si les filieres d'elimination sont saturees.",
               bold_prefix="Augmentation des capacites de stockage temporaire --")

    add_bullet(doc, " Negocier des clauses de montee en charge avec le prestataire de collecte DASRI "
               "(passage quotidien au lieu de 3 fois/semaine). Prevoir un prestataire secondaire "
               "en cas de defaillance.",
               bold_prefix="Renforcement des contrats d'enlevement --")

    add_bullet(doc, " Constituer un stock d'EPI (Equipements de Protection Individuelle) specifiques "
               "pour les agents de collecte. Renforcer les formations aux gestes de securite.",
               bold_prefix="Protection du personnel de collecte --")

    add_bullet(doc, " Renforcer les consignes de tri pour eviter les contaminations croisees. "
               "Assurer une tracabilite complete des DASRI (bordereau de suivi des dechets).",
               bold_prefix="Tri renforce et tracabilite --")

    # --- 1.3 HYGIENE ---
    doc.add_page_break()
    add_heading_styled(doc, "1.3 Hygiene et entretien des locaux", level=2)

    add_body(doc,
        "Le nettoyage et la desinfection des locaux sont des elements determinants dans la lutte "
        "contre les infections nosocomiales. En situation de crise sanitaire, les protocoles "
        "d'hygiene doivent etre intensifies, tant au niveau de la frequence que des produits utilises."
    )

    chart_path = gen_chart_hygiene(annual)
    doc.add_picture(chart_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    locaux_2015 = annual[
        (annual["indicateur"] == "Hygiène") &
        (annual["sous_indicateur"] == "Locaux") &
        (annual["year"] == 2015)
    ]
    vitres_2015 = annual[
        (annual["indicateur"] == "Hygiène") &
        (annual["sous_indicateur"] == "Vitres") &
        (annual["year"] == 2015)
    ]
    espaces_2015 = annual[
        (annual["indicateur"] == "Hygiène") &
        (annual["sous_indicateur"] == "Espaces verts") &
        (annual["year"] == 2015)
    ]

    if not locaux_2015.empty:
        loc = locaux_2015.iloc[0]
        vitres_val = vitres_2015.iloc[0]["value"] if not vitres_2015.empty else 0
        espaces_val = espaces_2015.iloc[0]["value"] if not espaces_2015.empty else 0
        add_constat_box(doc,
            f"En 2015, les surfaces de locaux a entretenir s'elevent a {loc['value']:,.0f} m² "
            f"en situation normale et augmenteraient a {loc['value_crise']:,.0f} m² en crise. "
            f"Les espaces verts ({espaces_val:,.0f} m²) et les vitres ({vitres_val:,.0f} m²) "
            f"sont egalement concernes."
        )

    # Profil mensuel hygiene
    add_body(doc,
        "Le profil mensuel de l'entretien des locaux montre la regularite de la charge "
        "de nettoyage au fil de l'annee :"
    )

    chart_path = gen_chart_hygiene_monthly(monthly)
    doc.add_picture(chart_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Propositions :")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLEU_FONCE

    add_bullet(doc, " Definir un protocole de desinfection de niveau crise avec des produits virucides "
               "et bactericides adaptes. Augmenter la frequence de nettoyage : passage de 1 a 3 fois/jour "
               "dans les zones de soins. Attention particuliere aux surfaces frequemment touchees "
               "(poignees, rampes, boutons d'ascenseur).",
               bold_prefix="Plan de bionettoyage renforce --")

    add_bullet(doc, " Maintenir un stock strategique de produits desinfectants pour couvrir 30 jours "
               "de consommation en mode crise. Diversifier les fournisseurs.",
               bold_prefix="Stocks de produits d'entretien --")

    add_bullet(doc, " Convention avec des entreprises de nettoyage pour un deploiement rapide d'equipes "
               "supplementaires (delai < 24h). Formation prealable aux protocoles hospitaliers.",
               bold_prefix="Mobilisation de personnel supplementaire --")

    add_bullet(doc, " Concentrer les efforts sur les zones de soins et les unites de reanimation. "
               "Reduire temporairement la frequence d'entretien des zones administratives.",
               bold_prefix="Priorisation des zones --")

    # --- 1.4 LINGERIE ---
    doc.add_page_break()
    add_heading_styled(doc, "1.4 Lingerie hospitaliere", level=2)

    add_body(doc,
        "Le linge hospitalier (draps, blouses, serviettes, tenues du personnel) est un element "
        "essentiel du confort des patients et de l'hygiene de l'etablissement. Un afflux de patients "
        "entraine une hausse directe des besoins en linge propre."
    )

    chart_path = gen_chart_lingerie(annual)
    doc.add_picture(chart_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    l2015 = annual[
        (annual["indicateur"] == "Lingerie") &
        (annual["sous_indicateur"] == "Quantité de linge") &
        (annual["year"] == 2015)
    ]

    if not l2015.empty:
        l = l2015.iloc[0]
        # Repartition par site
        plf_linge = annual_site[
            (annual_site["indicateur"] == "Lingerie") &
            (annual_site["sous_indicateur"] == "Quantité de linge") &
            (annual_site["year"] == 2015) &
            (annual_site["site_code"] == "PLF")
        ]["value"].sum()
        pct_plf_linge = (plf_linge / l["value"] * 100) if l["value"] > 0 else 0

        add_constat_box(doc,
            f"La quantite de linge traite en 2015 est de {l['value']:,.0f} kg en situation "
            f"normale. En crise, elle passerait a {l['value_crise']:,.0f} kg, soit pres de "
            f"{l['ecart']:,.0f} kg supplementaires par an. "
            f"Pitie-Salpetriere represente environ {pct_plf_linge:.0f} % du volume."
        )

    # Graphique par site
    chart_path = gen_chart_lingerie_site(annual_site)
    doc.add_picture(chart_path, width=Inches(5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Profil mensuel lingerie
    add_body(doc,
        "Le profil mensuel de la lingerie revele les variations saisonnieres dans le traitement "
        "du linge hospitalier :"
    )

    chart_path = gen_chart_lingerie_monthly(monthly)
    doc.add_picture(chart_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Propositions :")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLEU_FONCE

    add_bullet(doc, " Identifier une blanchisserie industrielle agree pouvant absorber le surplus. "
               "Tester cette convention une fois par an (exercice grandeur nature).",
               bold_prefix="Convention avec une blanchisserie de secours --")

    add_bullet(doc, " Constituer une reserve de draps et blouses a usage unique (non-tisses) pour "
               "les 15 premiers jours de crise. Ces articles limitent aussi le risque de transmission "
               "d'agents infectieux.",
               bold_prefix="Stock de linge a usage unique --")

    add_bullet(doc, " Augmenter la frequence de ramassage du linge souille (de 2 a 4 passages/jour). "
               "Prevoir des sacs hydrosolubles specifiques pour le linge contamine.",
               bold_prefix="Circuits de collecte renforces --")

    add_bullet(doc, " En crise, fournir des tenues a usage unique au personnel soignant. "
               "Prevoir des vestiaires temporaires aux entrees des unites d'isolement.",
               bold_prefix="Linge du personnel soignant --")

    # --- 1.5 MAGASIN ---
    doc.add_page_break()
    add_heading_styled(doc, "1.5 Magasin et approvisionnements", level=2)

    add_body(doc,
        "Le magasin central gere les fournitures hotelieres, les produits d'hygiene et les "
        "imprimes. En crise, la demande en fournitures de protection et d'hygiene augmente "
        "considerablement."
    )

    chart_path = gen_chart_magasin(annual)
    doc.add_picture(chart_path, width=Inches(5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    df_mag = annual[(annual["indicateur"] == "Magasin") & (annual["year"] == 2015)]
    total_n = df_mag["value"].sum()
    total_c = df_mag["value_crise"].sum()

    add_constat_box(doc,
        f"Le magasin gere environ {total_n:,.0f} references en situation normale. "
        f"En crise, le volume de commandes augmente de 70 % (soit {total_c:,.0f} references a gerer), "
        f"meme si le nombre de references catalogue reste stable."
    )

    # Profil mensuel magasin
    add_body(doc,
        "Le profil mensuel des approvisionnements permet d'anticiper les periodes de tension "
        "et d'adapter le calendrier de commandes :"
    )

    chart_path = gen_chart_magasin_monthly(monthly)
    doc.add_picture(chart_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Propositions :")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLEU_FONCE

    add_bullet(doc, " Maintenir en permanence un stock tampon de 30 jours pour les consommables "
               "critiques : gants, masques, surblouses, solutions hydroalcooliques, desinfectants. "
               "Systeme de rotation premier entre, premier sorti pour eviter la peremption.",
               bold_prefix="Stock strategique de 30 jours --")

    add_bullet(doc, " Pour chaque categorie de produit critique, avoir au minimum 2 fournisseurs "
               "references et une source alternative. Lecon du COVID-19 : eviter la dependance "
               "a un fournisseur unique.",
               bold_prefix="Diversification des fournisseurs --")

    add_bullet(doc, " Definir des seuils d'alerte pour chaque reference critique. Declenchement "
               "automatique d'une commande lorsqu'un stock passe sous le seuil.",
               bold_prefix="Systeme d'alerte de reapprovisionnement --")

    add_bullet(doc, " En cas de rupture, activer les circuits de depannage au sein du GHU AP-HP "
               "pour mutualiser les stocks disponibles entre etablissements.",
               bold_prefix="Coordination inter-etablissements --")

    # --- 1.6 VAGUEMESTRE ---
    doc.add_page_break()
    add_heading_styled(doc, "1.6 Service courrier (Vaguemestre)", level=2)

    add_body(doc,
        "Le service courrier assure la distribution du courrier aux patients, l'envoi des "
        "documents administratifs et la reception des colis. Bien que secondaire face aux enjeux "
        "sanitaires, la continuite du service courrier contribue au bien-etre des patients "
        "hospitalises et au bon fonctionnement administratif."
    )

    chart_path = gen_chart_vaguemestre(annual)
    doc.add_picture(chart_path, width=Inches(5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    plis = annual[
        (annual["indicateur"] == "Vaguemestre") &
        (annual["sous_indicateur"] == "Plis affranchis") &
        (annual["year"] == 2015)
    ]
    colis = annual[
        (annual["indicateur"] == "Vaguemestre") &
        (annual["sous_indicateur"] == "Colis") &
        (annual["year"] == 2015)
    ]

    if not plis.empty and not colis.empty:
        add_constat_box(doc,
            f"Plus de {plis.iloc[0]['value']:,.0f} plis affranchis par an et environ "
            f"{colis.iloc[0]['value']:,.0f} colis (2015). En crise, les volumes augmenteraient "
            f"significativement, notamment les envois administratifs."
        )

    # Profil mensuel vaguemestre
    add_body(doc,
        "Le profil mensuel du courrier revele les periodes de pointe et les creux saisonniers :"
    )

    chart_path = gen_chart_vaguemestre_monthly(monthly)
    doc.add_picture(chart_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Propositions :")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLEU_FONCE

    add_bullet(doc, " Privilegier l'envoi par e-mail ou via le portail patient pour les "
               "comptes-rendus et resultats. Reduire la dependance au courrier physique.",
               bold_prefix="Dematerialisation des courriers administratifs --")

    add_bullet(doc, " Prevoir un agent supplementaire mobilisable en cas de crise. "
               "Amenager les horaires de distribution pour les adapter aux contraintes de confinement.",
               bold_prefix="Renforcement ponctuel du service --")

    # --- SYNTHESE LOGISTIQUE ---
    doc.add_page_break()
    add_heading_styled(doc, "Synthese -- Logistique hospitaliere", level=2)

    chart_path = gen_chart_synthese(annual)
    doc.add_picture(chart_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Comparaison sites
    add_body(doc,
        "La comparaison entre les deux sites montre le poids relatif de chaque etablissement "
        "sur l'ensemble des indicateurs logistiques :"
    )

    chart_path = gen_chart_comparaison_sites(annual_site)
    doc.add_picture(chart_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tableau de synthese
    add_body(doc, "Tableau recapitulatif de l'impact d'une crise sanitaire sur la logistique (donnees 2015) :")

    df_2015 = annual[annual["year"] == 2015]
    synthese = df_2015.groupby("indicateur").agg({
        "value": "sum",
        "value_crise": "sum",
        "ecart": "sum"
    }).reset_index()

    headers = ["Domaine", "Volume normal", "Volume en crise", "Ecart", "Hausse"]
    rows = []
    for _, r in synthese.iterrows():
        rows.append([
            r["indicateur"],
            f"{r['value']:,.0f}",
            f"{r['value_crise']:,.0f}",
            f"+{r['ecart']:,.0f}",
            "+70 %"
        ])
    add_styled_table(doc, headers, rows)

    doc.add_paragraph()

    # Tableau par site
    add_body(doc, "Repartition par site (donnees 2015, situation normale) :")

    site_2015 = annual_site[annual_site["year"] == 2015]
    site_synth = site_2015.groupby(["indicateur", "site_code"])["value"].sum().unstack(fill_value=0).reset_index()

    headers_site = ["Domaine", "Pitie-Salpetriere (PLF)", "Charles Foix (CFX)", "Total"]
    rows_site = []
    for _, r in site_synth.iterrows():
        plf_v = r.get("PLF", 0)
        cfx_v = r.get("CFX", 0)
        rows_site.append([
            r["indicateur"],
            f"{plf_v:,.0f}",
            f"{cfx_v:,.0f}",
            f"{plf_v + cfx_v:,.0f}",
        ])
    add_styled_table(doc, headers_site, rows_site)

    doc.add_paragraph()

    # Plan d'action prioritaire
    add_heading_styled(doc, "Plan d'action prioritaire -- Logistique", level=3)

    headers_pa = ["Priorite", "Action", "Responsable"]
    rows_pa = [
        ["HAUTE", "Constituer un stock strategique d'EPI et de consommables (30 jours)", "Direction des Achats"],
        ["HAUTE", "Negocier des conventions avec prestataires de secours (restauration, linge, dechets)", "Direction Logistique"],
        ["HAUTE", "Definir un protocole de bionettoyage de niveau crise", "Service Hygiene"],
        ["HAUTE", "Augmenter les capacites de stockage DASRI", "Services Techniques"],
        ["MOYENNE", "Mettre en place un systeme d'alerte de reapprovisionnement automatique", "Pharmacie / Magasin"],
        ["MOYENNE", "Former le personnel logistique aux procedures de crise", "DRH / Formation Continue"],
        ["MOYENNE", "Prevoir des menus simplifies activables en 48h", "Service Restauration"],
        ["BASSE", "Dematerialiser le courrier administratif", "Direction des Systemes d'Information"],
        ["BASSE", "Reduire la frequence d'entretien des zones non-critiques en crise", "Direction Logistique"],
    ]
    add_styled_table(doc, headers_pa, rows_pa, col_widths=[2.5, 10, 4.5])

    # ============================
    # SECTIONS A VENIR (placeholders)
    # ============================
    doc.add_page_break()

    placeholder_sections = [
        ("2. Activite et services", "Volumes d'activite et organisation des soins"),
        ("3. Capacite d'accueil", "Lits, equipements et plateaux techniques"),
        ("4. Finance", "Equilibre budgetaire et investissements"),
        ("5. Patients", "Profil des patients et parcours de soins"),
        ("6. Qualite", "Satisfaction et indicateurs qualite"),
        ("7. Ressources humaines", "Effectifs et organisation du personnel"),
    ]

    for titre, desc in placeholder_sections:
        add_heading_styled(doc, titre, level=1)
        p = doc.add_paragraph()
        run = p.add_run(f"{desc} -- Section a completer.")
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = GRIS
        doc.add_paragraph()

    # ============================
    # SYNTHESE GENERALE
    # ============================
    doc.add_page_break()
    add_heading_styled(doc, "8. Synthese et recommandations", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Cette section sera completee une fois l'ensemble des categories analysees.")
    run.font.italic = True
    run.font.color.rgb = GRIS

    # --- Sauvegarde ---
    doc.save(OUTPUT_FILE)
    print(f"Rapport genere avec succes : {OUTPUT_FILE}")

    # Nettoyage des charts temporaires
    import shutil
    shutil.rmtree(CHARTS_DIR, ignore_errors=True)


if __name__ == "__main__":
    build_document()
